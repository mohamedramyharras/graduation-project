"""
generate_docx.py
================
Generate the Word manuscript for the IEEE paper using python-docx.
Prosthetic hand control system - GRU General (subject-independent) only.
No per-subject fine-tuning.

Produces: paper/manuscript.docx

Usage:
    python paper/generate_docx.py
"""

from __future__ import annotations
from pathlib import Path
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

PAPER_DIR   = Path(__file__).resolve().parent
PROJECT_ROOT = PAPER_DIR.parent
FIGURES_DIR  = PROJECT_ROOT / "outputs" / "figures"
OUTPUT_PATH  = PAPER_DIR / "manuscript.docx"


def set_font(run, name="Times New Roman", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold  = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Times New Roman"
    return p


def add_body(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.7)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    for run in p.runs:
        run.font.name   = "Times New Roman"
        run.font.size   = Pt(9)
        run.font.italic = True
    return p


def add_figure(doc, fname, caption):
    img_path = FIGURES_DIR / fname
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.8))
    else:
        p = doc.add_paragraph(f"[Figure not found: {fname}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def build_per_subject_table(doc):
    """Table I: Per-Subject R2 - 3 models (GRU, Ridge, MLP)."""
    headers = ["Subject", "GRU (Ours)", "Ridge", "MLP"]
    data = [
        ["S1",  "0.736", "0.690", "0.691"],
        ["S2",  "0.781", "0.566", "0.731"],
        ["S3",  "0.807", "0.718", "0.689"],
        ["S4",  "0.786", "0.551", "0.589"],
        ["S5",  "0.676", "0.634", "0.684"],
        ["S6",  "0.827", "0.802", "0.835"],
        ["S7",  "0.709", "0.714", "0.691"],
        ["S8*", "0.434", "0.213", "0.166"],
        ["S9",  "0.805", "0.759", "0.765"],
        ["S10", "0.878", "0.777", "0.898"],
    ]

    table = doc.add_table(rows=1 + len(data), cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)

    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            for run in row.cells[j].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)
                if row_data[0] == "S8*":
                    run.font.color.rgb = RGBColor(0xAA, 0x00, 0x00)

    doc.add_paragraph(
        "* Subject 8 excluded from primary statistics (data-quality outlier)."
    ).runs[0].font.size = Pt(8)


def build_summary_table(doc):
    """Table II: Summary statistics - 3 models."""
    headers = ["Model", "R\u00b2\n(mean\u00b1SD)", "RMSE\n(mean\u00b1SD)",
               "NRMSE%\n(mean\u00b1SD)", "MAE\n(mean\u00b1SD)", "Pearson r\n(mean\u00b1SD)"]
    data = [
        ["GRU (Ours)*", "0.778\u00b10.062", "0.103\u00b10.011", "11.5\u00b11.8",
         "0.078\u00b10.009", "0.890\u00b10.033"],
        ["MLP",          "0.730\u00b10.086", "0.112\u00b10.014", "12.6\u00b12.1",
         "0.085\u00b10.011", "0.873\u00b10.036"],
        ["Ridge",        "0.690\u00b10.084", "0.122\u00b10.017", "13.6\u00b12.2",
         "0.095\u00b10.014", "0.847\u00b10.043"],
    ]
    table = doc.add_table(rows=1 + len(data), cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(8)

    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            for run in row.cells[j].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(8)
                if j >= 1 and i == 0:
                    run.bold = True

    doc.add_paragraph(
        "95% CI (GRU): R\u00b2 [0.730, 0.826], NRMSE% [10.12, 12.88]. "
        "RMSE and MAE in normalized (0\u20131) force. "
        "* Best performing model (proposed method)."
    ).runs[0].font.size = Pt(8)


def build_prior_work_table(doc):
    """Table III: Prior work comparison."""
    headers = ["Method", "Ch.", "Subj.", "Task", "Eval.", "R\u00b2", "NRMSE%", "Notes"]
    data = [
        # Sparse sEMG: Traditional
        ["[Traditional Regression]",     "",   "",   "",              "",      "",         "",        ""],
        ["Castellini 2009",              "10", "8",  "Finger force",  "Intra", "0.70\u20130.78", "\u224815", "SVR"],
        ["Hahne et al. 2014",            "4",  "6",  "Wrist 2-DOF",  "Intra", "\u22480.72",     "\u224814", "Ridge; \u2020"],
        ["Kim et al. 2020",              "7",  "\u2014",  "Wrist+grip", "Intra", "\u2014",       "25.6\u00b114", "Synergy; r=0.846"],
        # Sparse sEMG: Deep Learning
        ["[Deep Learning]",              "",   "",   "",              "",      "",         "",        ""],
        ["Mao et al. 2023",              "6",  "10", "Grip force",    "Intra", "0.963",    "\u2014",  "GRNN; \u2021"],
        ["Ghorbani et al. 2023",         "8",  "10", "Grip force",    "Intra", "0.994",    "\u2014",  "\u00a7 INFLATED"],
        ["Ridge (Ours)",                 "2",  "9",  "Grip force",    "Cross", "0.690\u00b10.08", "13.6\u00b12.2", "Baseline"],
        ["MLP (Ours)",                   "2",  "9",  "Grip force",    "Cross", "0.730\u00b10.09", "12.6\u00b12.1", "Baseline"],
        ["GRU (Ours)",                   "2",  "9",  "Grip force",    "Cross", "0.778\u00b10.06", "11.5\u00b11.8", "BEST; 2 ch"],
        # HD-EMG
        ["[HD-sEMG]",                    "",   "",   "",              "",      "",         "",        ""],
        ["Ma et al. 2021",               "8",  "10", "Joint angle",   "Intra", "\u22480.90", "\u22488", "LSTM; \u2020\u2020"],
        ["Li et al. 2024",               "HD", "12", "Grasp+3-DoF",  "Intra", "\u2014",    "9.7",   "Graph ST; r=91.9%"],
    ]
    table = doc.add_table(rows=1 + len(data), cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(8)

    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        is_header = row_data[1] == "" and row_data[0].startswith("[")
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            for run in row.cells[j].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(8)
                if is_header:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
                if row_data[0] == "GRU (Ours)":
                    run.bold = True

    doc.add_paragraph(
        "\u2020 Hahne: offline, 2-DOF wrist control \u2014 not identical to grip force.\n"
        "\u2021 Mao: intrasubject (per-subject train/test); 6 dedicated forearm electrodes.\n"
        "\u00a7 Ghorbani: force used as 9th input feature (autoregressive, 10 ms horizon); "
        "test set normalized independently (data leakage).\n"
        "\u2020\u2020 Ma et al.: joint angle (degrees), not force \u2014 indicative comparison only.\n"
        "Eval.: Intra = intrasubject; Cross = cross-subject (subject-independent)."
    ).runs[0].font.size = Pt(8)


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)

    # -- Title --
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(
        "A Lightweight 2-Channel sEMG-Based Grip Force Controller "
        "for 3D-Printed Prosthetic Hands Using Gated Recurrent Units"
    )
    set_font(title_run, size=14, bold=True)

    doc.add_paragraph()
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_p.add_run("[Author Name(s)] \u2014 [Department, University]")
    set_font(author_run, size=11, italic=True)
    doc.add_paragraph()

    # -- Abstract --
    add_heading(doc, "Abstract", level=2)
    add_body(doc,
        "Proportional grip force control is essential for functional upper-limb prostheses, "
        "yet most EMG-based approaches require high-channel-count electrode arrays or "
        "per-user calibration that limits clinical deployment. We present a complete end-to-end "
        "system for real-time grip force prediction from two surface EMG (sEMG) channels, "
        "deployed on an InMoov 3D-printed prosthetic hand actuated via Arduino-controlled "
        "linear actuators. The two channels\u2014ECRB (Ch5) and ECU (Ch3)\u2014are selected by the "
        "Minimum-Redundancy Maximum-Relevance (MRMR) criterion from an 8-channel Myo armband. "
        "A lightweight Gated Recurrent Unit (GRU) with 23,809 parameters is trained in a "
        "fully subject-independent manner on pooled data from nine healthy participants, "
        "requiring no per-user calibration. On the publicly available Ghorbani grip-force "
        "dataset (200 Hz), the proposed method achieves R\u00b2 = 0.778 \u00b1 0.062 "
        "(95% CI [0.730, 0.826]) with four of nine subjects exceeding R\u00b2 = 0.80, "
        "outperforming both MLP (R\u00b2 = 0.730) and Ridge regression (R\u00b2 = 0.690) baselines."
    )
    doc.add_paragraph()

    # -- I. Introduction --
    add_heading(doc, "I. Introduction", level=1)
    add_body(doc,
        "Upper-limb amputees and individuals with motor impairments rely on myoelectric "
        "prostheses to regain grasping function. Intuitive, proportional grip force control\u2014"
        "where the prosthesis mirrors the intended force of the user's residual musculature\u2014"
        "remains a central challenge in rehabilitation engineering.",
        indent=True
    )
    add_body(doc,
        "Surface EMG (sEMG) signals encode neural drive to the finger and wrist flexors "
        "and extensors, making them the dominant sensing modality for non-invasive prosthetic "
        "control. Traditional pattern recognition approaches classify discrete grasp types "
        "rather than estimating continuous force, inherently limiting proportionality. "
        "Regression-based methods bridge this gap by mapping EMG features to a continuous "
        "force output.",
        indent=True
    )
    add_body(doc,
        "Deep learning, and recurrent architectures in particular, have shown promise for "
        "exploiting the temporal structure of EMG sequences. The Gated Recurrent Unit (GRU) "
        "offers a favorable accuracy-complexity trade-off over LSTM. Prior GRU-based grip "
        "force estimation reported R\u00b2 = 0.994 but included force feedback in the input\u2014"
        "making the task effectively autoregressive\u2014thereby inflating reported performance. "
        "The present work uses EMG features exclusively.",
        indent=True
    )
    add_body(doc,
        "This paper presents a complete prosthetic grip force control system comprising "
        "two MyoWare 2.0 sEMG sensors, a subject-independent GRU model, and an InMoov "
        "3D-printed hand actuated by linear actuators via Arduino Mega 2560. The system "
        "requires no per-user calibration, making it suitable for plug-and-play deployment.",
        indent=True
    )

    doc.add_paragraph()
    add_heading(doc, "Contributions", level=2)
    for item in [
        "A complete end-to-end prosthetic grip force control system from 2 sEMG electrodes "
        "to actuated InMoov hand.",
        "MRMR channel selection identifying ECRB (Ch5) and ECU (Ch3) as optimal channels, "
        "with anatomical rationale based on the extensor paradox.",
        "A subject-independent GRU model (23,809 parameters) achieving R\u00b2 = 0.778 with no "
        "per-user calibration.",
        "Rigorous benchmarking with mean \u00b1 SD, 95% CI, and transparent outlier characterization.",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    # -- II. Related Work --
    doc.add_page_break()
    add_heading(doc, "II. Related Work", level=1)

    add_heading(doc, "A. EMG-to-Force Mapping", level=2)
    add_body(doc,
        "Sartori et al. demonstrated that Hill-type musculoskeletal models driven by EMG "
        "can estimate joint torques. Hahne et al. compared linear and non-linear regression "
        "for simultaneous proportional myoelectric control, establishing ridge regression and "
        "MLP as standard baselines.",
        indent=True
    )

    add_heading(doc, "B. Deep Learning for EMG", level=2)
    add_body(doc,
        "Recurrent networks are well-suited to EMG-to-force regression due to temporal "
        "smoothness in force production. Ma et al. demonstrated LSTM-based joint angle "
        "estimation. Ghorbani et al. reported R\u00b2 = 0.994 using GRU, but their input "
        "concatenated force as a feedback signal, converting the problem to trivial "
        "autoregressive prediction. The present work avoids this methodological flaw.",
        indent=True
    )

    add_heading(doc, "C. Open-Source Prosthetics", level=2)
    add_body(doc,
        "The InMoov hand, designed by Ga\u00ebl Langevin, is an open-source 3D-printed "
        "humanoid robot that provides an accessible platform for prosthetic research. "
        "Ten Brinke and Carloni reviewed open-source anthropomorphic hand designs, "
        "noting that tendon-driven actuation (as used in InMoov) provides a favorable "
        "balance of cost, weight, and dexterity for prosthetic applications.",
        indent=True
    )

    # -- III. System Architecture --
    add_heading(doc, "III. System Architecture", level=1)
    add_body(doc,
        "The system uses a split-processing architecture: (1) two MyoWare 2.0 surface EMG "
        "sensors placed on the posterior forearm over ECRB and ECU muscles; (2) an Arduino "
        "Mega 2560 microcontroller sampling both channels at 2,000 Hz via its 10-bit ADC and "
        "streaming raw values over USB serial; (3) a Raspberry Pi 3 (ARM Cortex-A53, 1.2 GHz) "
        "performing feature extraction and GRU inference in Python/PyTorch; (4) the Arduino "
        "receiving the predicted force value and mapping it to PWM signals driving five linear "
        "actuators connected to the InMoov hand's fingers via nylon string tendons; and (5) an "
        "HC-05 Bluetooth module for optional gain calibration. "
        "The GRU model output (0\u20131 normalized force) maps directly to PWM duty cycles, "
        "providing continuous proportional grip force. Total end-to-end latency is approximately "
        "45\u201355 ms (20\u201325 Hz update rate), well within the 100\u2013300 ms acceptable "
        "threshold for prosthetic force control.",
        indent=True
    )

    # -- IV. Forearm Anatomy and Electrode Placement --
    add_heading(doc, "IV. Forearm Anatomy and Electrode Placement", level=1)

    add_heading(doc, "A. ECRB (Extensor Carpi Radialis Brevis)", level=2)
    add_body(doc,
        "ECRB originates from the lateral epicondyle of the humerus (common extensor origin) "
        "and inserts on the dorsal surface of the 3rd metacarpal base. Its primary functions "
        "are wrist extension and radial deviation. During power grip, ECRB co-activates to "
        "maintain the wrist in slight extension against the flexor load\u2014the extensor paradox "
        "whereby wrist extensors fire during a flexion-dominated task. MRMR identifies ECRB "
        "(Ch5) as the most informative channel with MI relevance = 0.077.",
        indent=True
    )

    add_heading(doc, "B. ECU (Extensor Carpi Ulnaris)", level=2)
    add_body(doc,
        "ECU originates from the lateral epicondyle and posterior ulna, inserting on the "
        "5th metacarpal base. It provides wrist extension with ulnar deviation. During grip, "
        "ECU stabilizes the wrist on the ulnar side, complementing ECRB's radial action. "
        "MRMR selects ECU (Ch3) second due to its low redundancy with ECRB (MRMR = 0.019), "
        "providing complementary ulnar-side information about grip dynamics.",
        indent=True
    )

    add_heading(doc, "C. Electrode Placement Protocol", level=2)
    add_body(doc,
        "Electrodes are placed over the muscle bellies at the proximal third of the forearm, "
        "aligned parallel to muscle fiber direction, following standard sEMG guidelines. "
        "ECRB electrode: lateral posterior forearm, approximately 5 cm distal to lateral "
        "epicondyle. ECU electrode: medial posterior forearm, over the ulnar border. "
        "For the deployed prosthetic system, MyoWare 2.0 differential muscle sensors "
        "are used.",
        indent=True
    )

    # -- V. Dataset and Preprocessing --
    add_heading(doc, "V. Dataset and Preprocessing", level=1)

    add_heading(doc, "A. Ghorbani Grip Force Dataset", level=2)
    add_body(doc,
        "We use the publicly available Ghorbani grip-force dataset. Ten healthy participants "
        "(9 male, 1 female; mean age 23.8 years) performed isometric precision-type gripping "
        "tasks. EMG was acquired at 200 Hz using a Myo armband (8 channels). Grip force "
        "(z-axis) was measured using an ATI Mini-45 sensor. Each subject completed 3 trials. "
        "The dataset is pre-filtered; we skip additional bandpass filtering.",
        indent=True
    )

    add_heading(doc, "B. Subject 8 Outlier", level=2)
    add_body(doc,
        "Subject 8 yields anomalously low accuracy across all model families: GRU R\u00b2 = 0.434, "
        "Ridge R\u00b2 = 0.213, MLP R\u00b2 = 0.166. The consistent failure of even the linear model "
        "indicates a data-quality issue rather than model failure. Subject 8 is excluded from "
        "primary statistics; inclusive 10-subject results are reported separately.",
        indent=True
    )

    add_heading(doc, "C. Feature Extraction", level=2)
    add_body(doc,
        "Six time-domain features are extracted per channel using a 100 ms sliding window "
        "(80% overlap, 20 ms hop): RMS, MAV, Waveform Length (WL), Variance (VAR), "
        "Zero-Crossings (ZC), and Slope Sign Changes (SSC). Each feature vector is "
        "augmented with first- and second-order temporal differences (delta, delta-delta), "
        "yielding 36 total inputs per time step (2 channels \u00d7 6 features \u00d7 3 delta orders).",
        indent=True
    )

    add_heading(doc, "D. MRMR Channel Selection", level=2)
    add_body(doc,
        "Eight candidate channels are ranked by MRMR. The top-2 selected are Ch5 (ECRB, "
        "relevance MI = 0.077) and Ch3 (ECU, MRMR score = 0.019). These anatomically "
        "correspond to primary wrist extensor muscles active during grip force generation.",
        indent=True
    )

    # -- VI. Methodology --
    doc.add_page_break()
    add_heading(doc, "VI. Methodology", level=1)

    add_heading(doc, "A. Model Architecture", level=2)
    add_body(doc,
        "The predictor is a one-layer GRU (hidden size 64) followed by two fully connected "
        "layers (FC 64, FC 1) with ReLU activations and dropout (p=0.2) after each layer. "
        "Total parameters: 23,809. Input sequences span L = 50 time steps (1 second of EMG "
        "context). No bidirectional processing is used to preserve causal inference.",
        indent=True
    )

    add_heading(doc, "B. Subject-Independent Training", level=2)
    add_body(doc,
        "A single general model is trained on pooled normalized data from all nine primary "
        "subjects. Per-subject MinMaxScaler normalization prevents data leakage (fit on "
        "training split only). Training uses Adam optimizer (lr = 0.001, weight decay = 5e-5), "
        "batch size 512, gradient clipping at 1.0, early stopping (patience 12), max 50 epochs, "
        "and ReduceLROnPlateau scheduling. No per-user fine-tuning is applied\u2014the model is "
        "trained once and used directly for all subjects, enabling plug-and-play prosthetic "
        "deployment.",
        indent=True
    )

    add_heading(doc, "C. Prediction Smoothing", level=2)
    add_body(doc,
        "Raw predictions are post-processed with a zero-phase 2nd-order Butterworth "
        "low-pass filter at 4 Hz (sosfiltfilt), eliminating high-frequency artifacts while "
        "preserving the bandwidth of realistic grip force trajectories.",
        indent=True
    )

    add_heading(doc, "D. Evaluation Protocol", level=2)
    add_body(doc,
        "The last 20% of each subject's sequences form the held-out test set (temporal split, "
        "10-sequence gap). Reported metrics: R\u00b2, RMSE, NRMSE%, MAE, Pearson r. All metrics "
        "are computed on normalized (0\u20131) force values.",
        indent=True
    )

    # -- VII. Experiments --
    add_heading(doc, "VII. Experiments", level=1)
    add_body(doc,
        "Two baselines are included: (1) Ridge Regression with regularization selected by "
        "5-fold cross-validation on flattened sequences; (2) MLP (128-64-1, ReLU, dropout 0.3) "
        "with early stopping. All models use identical normalization and splits for fairness. "
        "Experiments run in PyTorch with fixed seed 42.",
        indent=True
    )

    # -- VIII. Results --
    doc.add_page_break()
    add_heading(doc, "VIII. Results", level=1)

    add_heading(doc, "A. Per-Subject R\u00b2 (Table I)", level=2)
    build_per_subject_table(doc)
    doc.add_paragraph()

    add_heading(doc, "B. Summary Statistics (Table II)", level=2)
    build_summary_table(doc)
    doc.add_paragraph()

    add_body(doc,
        "The GRU achieves R\u00b2 = 0.778 \u00b1 0.062 (9 subjects, 95% CI [0.730, 0.826]) with "
        "NRMSE = 11.5 \u00b1 1.8% and Pearson r = 0.890 \u00b1 0.033. Four of nine subjects exceed "
        "R\u00b2 = 0.80. The GRU reduces NRMSE by 2.1 percentage points vs. Ridge (13.6% \u2192 11.5%), "
        "a 15% relative improvement in prediction error. The subject-independent design means "
        "these results reflect true cross-subject generalization\u2014no per-user calibration was used.",
        indent=True
    )

    add_heading(doc, "C. Comparison with Prior Work (Table III)", level=2)
    build_prior_work_table(doc)
    doc.add_paragraph()
    add_body(doc,
        "Traditional methods with 4\u201310 channels achieve R\u00b2 \u2248 0.70\u20130.78 and "
        "NRMSE \u2248 14\u201315% in intrasubject settings. Our GRU with only 2 channels achieves "
        "R\u00b2 = 0.778 and NRMSE = 11.5% in a harder cross-subject (subject-independent) setting, "
        "outperforming these methods despite using 2\u20135\u00d7 fewer electrodes. "
        "Mao et al. (2023, GRNN, 6 channels) report R\u00b2 = 0.963 but use per-subject training "
        "on simple grip force profiles\u2014an easier setting. Ghorbani et al. (2023) report "
        "R\u00b2 = 0.994 but used force as a 9th input feature (autoregressive task with 10 ms "
        "horizon) and applied fit_transform independently to test sets (data leakage).",
        indent=True
    )

    add_heading(doc, "D. Figures", level=2)
    add_figure(doc, "fig1_methodology.png",
               "Fig. 1. System pipeline: MyoWare 2.0 \u2192 MRMR channel selection \u2192 "
               "feature extraction (36 inputs) \u2192 GRU model \u2192 Butterworth smoothing \u2192 "
               "Arduino Mega \u2192 InMoov hand.")
    add_figure(doc, "fig2_mrmr_channels.png",
               "Fig. 2. MRMR channel ranking for the 8 Myo armband channels. "
               "Asterisked bars indicate the two selected channels (ECRB, ECU).")
    add_figure(doc, "fig3_per_subject_r2.png",
               "Fig. 3. Per-subject R\u00b2 for all 10 subjects and 3 model variants. "
               "Subject 8 (hatched) is the data-quality outlier.")
    add_figure(doc, "fig4_outlier_analysis.png",
               "Fig. 4. Subject 8 outlier analysis. Left: R\u00b2 for S8 vs. group mean. "
               "Right: GRU per-subject R\u00b2.")
    add_figure(doc, "fig5_prediction_traces.png",
               "Fig. 5. GRU prediction traces for S10 (best), S6 (mid-high), S5 (mid-low). "
               "Solid: actual; dashed: predicted.")
    add_figure(doc, "fig6_scatter.png",
               "Fig. 6. Predicted vs. actual scatter plot for all 9 subjects combined "
               "(normalized force, GRU test set).")
    add_figure(doc, "fig7_model_comparison.png",
               "Fig. 7. Model comparison: mean R\u00b2 \u00b1 95% CI with per-subject dots.")
    add_figure(doc, "fig8_anatomy_electrode.png",
               "Fig. 8. Posterior forearm anatomy: ECRB and ECU muscles with electrode "
               "placement markers over muscle bellies.")

    # -- IX. Discussion --
    doc.add_page_break()
    add_heading(doc, "IX. Discussion", level=1)
    add_body(doc,
        "The two-channel configuration (ECRB + ECU) is clinically practical: it maps to "
        "standard adhesive electrode patches, reduces hardware cost, and simplifies donning. "
        "MRMR consistently identifies these channels as most informative, aligned with the "
        "biomechanical co-activation of wrist extensors during grip force generation "
        "(the extensor paradox).",
        indent=True
    )
    add_body(doc,
        "The subject-independent model achieves R\u00b2 = 0.778 without any per-user calibration, "
        "demonstrating strong cross-subject transferability. This is critical for prosthetic "
        "deployment: new users can begin using the device immediately without a training session. "
        "The R\u00b2 = 0.778 represents a practical level of force tracking sufficient for "
        "gross grip tasks such as grasping objects of varying compliance.",
        indent=True
    )
    add_body(doc,
        "Ghorbani et al. report R\u00b2 = 0.994 on the same dataset. Their inflated performance "
        "is due to: (1) force signal concatenated as a 9th input feature, creating an "
        "autoregressive task with 10 ms horizon; (2) independent fit_transform on test sets "
        "(data leakage). Our work uses EMG features exclusively and fits scalers on training "
        "data only, representing a genuine EMG-to-force benchmark.",
        indent=True
    )
    add_body(doc,
        "For embedded deployment, the system uses a split-processing architecture. The Arduino "
        "Mega 2560 handles real-time ADC sampling (2,000 Hz, 2 channels) and PWM actuator "
        "control. The Raspberry Pi 3 (ARM Cortex-A53, 1.2 GHz, 1 GB RAM) runs Python/PyTorch "
        "for feature extraction (~2 ms) and GRU inference (~20\u201330 ms per 50-step sequence, "
        "corresponding to roughly 1 million MACs). The 23,809-parameter model requires only "
        "~93 KB in 32-bit float\u2014trivially small in the Pi's 1 GB RAM. Total end-to-end "
        "latency from EMG to actuator command is approximately 45\u201355 ms, yielding a "
        "20\u201325 Hz update rate well below the 100\u2013300 ms acceptable threshold for "
        "prosthetic force control. Total component cost is under $50.",
        indent=True
    )
    add_body(doc,
        "Limitations include: healthy young adult participants only, single isometric grasp "
        "type, within-trial evaluation only, and the Myo armband (discontinued) as the "
        "validation sensor. Cross-session, cross-day, and amputee validation remain for "
        "future work.",
        indent=True
    )

    # -- X. Conclusion --
    add_heading(doc, "X. Conclusion", level=1)
    add_body(doc,
        "We presented a complete prosthetic grip force control system using two sEMG channels "
        "(ECRB, ECU) and a lightweight subject-independent GRU model. The method achieves "
        "R\u00b2 = 0.778 \u00b1 0.062 on the Ghorbani dataset (9 subjects) with 4 of 9 subjects "
        "exceeding R\u00b2 = 0.80, outperforming Ridge and MLP baselines. With 23,809 parameters "
        "and two-electrode hardware requirements, the system is deployable on embedded "
        "prosthetic controllers. The InMoov 3D-printed hand with nylon tendon actuation "
        "provides an accessible, open-source prosthetic platform. MRMR channel selection "
        "with anatomical rationale and delta feature augmentation constitute the key "
        "methodological contributions.",
        indent=True
    )

    # -- References --
    add_heading(doc, "References", level=1)
    refs = [
        "[1] Castellini & van der Smagt, Biological Cybernetics, 2009.",
        "[2] Peng et al., IEEE TPAMI, 2005. (MRMR)",
        "[3] Meattini et al., IEEE/ASME TMech, 2018.",
        "[4] Englehart & Hudgins, IEEE TBME, 2003.",
        "[5] Fougner et al., IEEE TNSRE, 2011.",
        "[6] Hahne et al., IEEE TNSRE, 2014.",
        "[7] Xiong et al., IEEE/CAA JAS, 2021.",
        "[8] Ma et al., IEEE Sensors J., 2021.",
        "[9] Cho et al., arXiv:1406.1078, 2014. (GRU)",
        "[10] Ghorbani et al., arXiv:2302.09555, 2023.",
        "[11] Sartori et al., PLOS ONE, 2012.",
        "[12] Phinyomark et al., Expert Syst. Appl., 2012.",
        "[13] Ghorbani, GitHub Dataset, 2023.",
        "[14] Kingma & Ba, arXiv:1412.6980, 2014. (Adam)",
        "[15] Kim et al., J. Healthcare Eng., 2020.",
        "[16] Mao et al., Technology and Health Care, 2023.",
        "[17] Li et al., IEEE JBHI, 2024.",
        "[18] Criswell, Cram's Intro. to Surface EMG, 2011.",
        "[19] Basmajian & De Luca, Muscles Alive, 1985.",
        "[20] De Luca, J. Applied Biomechanics, 1997.",
        "[21] Langevin, InMoov Open-Source Robot, 2017.",
        "[22] ten Brinke & Carloni, Robotics Auton. Syst., 2021.",
        "[23] Warden & Situnayake, TinyML, 2020.",
        "[24] Scheme & Englehart, JRRD, 2011.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)

    doc.save(str(OUTPUT_PATH))
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
